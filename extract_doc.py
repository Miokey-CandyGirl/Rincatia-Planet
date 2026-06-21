import sys
sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document
    doc = Document('consultation/现代琳凯蒂亚语 v26.5.docx')
    for p in doc.paragraphs:
        print(p.text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            print(' | '.join(cells))
except ImportError:
    print("python-docx not installed")
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        with zipfile.ZipFile('consultation/现代琳凯蒂亚语 v26.5.docx', 'r') as z:
            xml_content = z.read('word/document.xml')
            tree = ET.fromstring(xml_content)
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            for t in tree.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                print(t.text, end='')
    except Exception as e:
        print(f"Error: {e}")