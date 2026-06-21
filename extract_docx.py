import docx
import os

doc_path = r'c:\Users\CandyGirl\Desktop\光线传奇\琳凯蒂亚语\琳凯蒂亚文化网站设计2026.06\consultation\现代琳凯蒂亚语 v26.5.docx'
doc = docx.Document(doc_path)

with open('docx_content.txt', 'w', encoding='utf-8') as f:
    for i, p in enumerate(doc.paragraphs):
        f.write(f'[{i}] {p.text}\n')

print(f'Total paragraphs: {len(doc.paragraphs)}')
print('Content saved to docx_content.txt')